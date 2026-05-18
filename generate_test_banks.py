#!/usr/bin/env python3
"""生成测试题库文件"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_DIR = "/sessions/6a0ab8a9e859ff74c089ce32/workspace/quiz-app/test-banks"
os.makedirs(OUTPUT_DIR, exist_ok=True)


def add_question(doc, num, text, options, answer, explanation=""):
    """添加一道题"""
    p = doc.add_paragraph()
    p.add_run(f"{num}. ").bold = True
    p.add_run(text)

    for opt in options:
        doc.add_paragraph(opt)

    p = doc.add_paragraph()
    p.add_run(f"答案：{answer}").bold = True

    if explanation:
        p = doc.add_paragraph()
        p.add_run(f"解析：{explanation}")

    doc.add_paragraph("---")


def create_computer_basics():
    """计算机基础知识题库"""
    doc = Document()

    # 单选题
    add_question(doc, 1,
        "计算机中，1KB等于多少字节？",
        ["A. 1000字节", "B. 1024字节", "C. 100字节", "D. 512字节"],
        "B",
        "1KB = 1024字节，这是计算机存储的基本换算单位。"
    )

    add_question(doc, 2,
        "以下哪个不是计算机的输入设备？",
        ["A. 键盘", "B. 鼠标", "C. 显示器", "D. 扫描仪"],
        "C",
        "显示器是输出设备，用于显示计算机处理后的信息。键盘、鼠标和扫描仪都是输入设备。"
    )

    add_question(doc, 3,
        "HTTP协议的默认端口号是？",
        ["A. 21", "B. 22", "C. 80", "D. 443"],
        "C",
        "HTTP默认端口是80，HTTPS默认端口是443，FTP默认端口是21，SSH默认端口是22。"
    )

    add_question(doc, 4,
        "在Windows系统中，用于查看IP地址的命令是？",
        ["A. ping", "B. ipconfig", "C. tracert", "D. netstat"],
        "B",
        "ipconfig用于查看和配置Windows系统的IP地址。ping用于测试网络连通性。"
    )

    add_question(doc, 5,
        "以下哪种文件格式是图片文件？",
        ["A. .mp3", "B. .docx", "C. .png", "D. .exe"],
        "C",
        ".png是图片格式，.mp3是音频格式，.docx是Word文档格式，.exe是可执行文件格式。"
    )

    add_question(doc, 6,
        "CPU的中文名称是？",
        ["A. 内存", "B. 硬盘", "C. 中央处理器", "D. 显卡"],
        "C",
        "CPU（Central Processing Unit）即中央处理器，是计算机的核心运算部件。"
    )

    add_question(doc, 7,
        "以下哪个不是操作系统？",
        ["A. Windows", "B. Linux", "C. Chrome", "D. macOS"],
        "C",
        "Chrome是浏览器，不是操作系统。Windows、Linux和macOS都是操作系统。"
    )

    add_question(doc, 8,
        "计算机病毒的本质是？",
        ["A. 生物病毒", "B. 硬件故障", "C. 一段程序代码", "D. 电磁干扰"],
        "C",
        "计算机病毒本质上是一段具有自我复制能力的恶意程序代码。"
    )

    # 多选题
    add_question(doc, 9,
        "以下哪些是常见的办公软件？（多选）",
        ["A. Microsoft Word", "B. Adobe Photoshop", "C. WPS Office", "D. Microsoft Excel", "E. Microsoft PowerPoint"],
        "ACDE",
        "Word、WPS Office、Excel和PowerPoint都是办公软件。Photoshop是图像处理软件。"
    )

    add_question(doc, 10,
        "以下哪些属于计算机网络协议？（多选）",
        ["A. TCP/IP", "B. HTTP", "C. FTP", "D. CPU"],
        "ABC",
        "TCP/IP、HTTP和FTP都是网络协议。CPU是中央处理器，不是协议。"
    )

    # 判断题
    add_question(doc, 11,
        "计算机断电后，硬盘中的数据会丢失。",
        ["A. 正确", "B. 错误"],
        "错误",
        "硬盘是外部存储设备，断电后数据不会丢失。只有内存（RAM）中的数据在断电后会丢失。"
    )

    add_question(doc, 12,
        "IPv4地址由32位二进制数组成。",
        ["A. 正确", "B. 错误"],
        "正确",
        "IPv4地址由32位二进制数组成，通常用点分十进制表示，如192.168.1.1。"
    )

    add_question(doc, 13,
        "Wi-Fi是一种有线网络连接方式。",
        ["A. 正确", "B. 错误"],
        "错误",
        "Wi-Fi是一种无线网络连接技术，不需要网线即可连接网络。"
    )

    add_question(doc, 14,
        "防火墙可以完全防止所有网络攻击。",
        ["A. 正确", "B. 错误"],
        "错误",
        "防火墙可以阻挡大部分外部攻击，但不能完全防止所有网络攻击，需要配合其他安全措施。"
    )

    add_question(doc, 15,
        "一个字节（Byte）等于8个比特（bit）。",
        ["A. 正确", "B. 错误"],
        "正确",
        "1 Byte = 8 bit，这是计算机存储的基本单位换算关系。"
    )

    doc.save(os.path.join(OUTPUT_DIR, "计算机基础题库.docx"))
    print(f"✓ 计算机基础题库.docx (15题: 8单选+2多选+5判断)")


def create_english_basics():
    """英语基础题库"""
    doc = Document()

    add_question(doc, 1,
        "What is the past tense of 'go'?",
        ["A. goed", "B. went", "C. gone", "D. going"],
        "B",
        "'go' 的过去式是 'went'，这是不规则动词变化。"
    )

    add_question(doc, 2,
        "Choose the correct sentence:",
        ["A. She don't like coffee.", "B. She doesn't likes coffee.", "C. She doesn't like coffee.", "D. She not like coffee."],
        "C",
        "第三人称单数否定形式：主语 + doesn't + 动词原形。"
    )

    add_question(doc, 3,
        "The word 'beautiful' is a(n) ___.",
        ["A. noun", "B. verb", "C. adjective", "D. adverb"],
        "C",
        "'beautiful' 意为'美丽的'，是形容词(adjective)。"
    )

    add_question(doc, 4,
        "Which word means '快速地'?",
        ["A. slowly", "B. quickly", "C. quietly", "D. loudly"],
        "B",
        "'quickly' 意为'快速地'，'slowly' 意为'慢慢地'。"
    )

    add_question(doc, 5,
        "Fill in the blank: I ___ to school every day.",
        ["A. go", "B. goes", "C. going", "D. went"],
        "A",
        "主语是 'I'，一般现在时动词用原形 'go'。"
    )

    add_question(doc, 6,
        "What is the opposite of 'hot'?",
        ["A. warm", "B. cool", "C. cold", "D. wet"],
        "C",
        "'hot' 的反义词是 'cold'（冷）。"
    )

    add_question(doc, 7,
        "Choose the correct preposition: The book is ___ the table.",
        ["A. in", "B. on", "C. at", "D. to"],
        "B",
        "书在桌子上用介词 'on'，表示在某物的表面上。"
    )

    add_question(doc, 8,
        "Which sentence is in the present continuous tense?",
        ["A. I eat breakfast.", "B. I ate breakfast.", "C. I am eating breakfast.", "D. I will eat breakfast."],
        "C",
        "'I am eating breakfast.' 使用了 be + doing 结构，是现在进行时。"
    )

    add_question(doc, 9,
        "How many days are there in a week?",
        ["A. 5", "B. 6", "C. 7", "D. 8"],
        "C",
        "一周有7天：Monday到Sunday。"
    )

    add_question(doc, 10,
        "Which is correct?",
        ["A. He is more taller than me.", "B. He is taller than me.", "C. He is tallest than me.", "D. He is tall than me."],
        "B",
        "比较级：taller + than。不需要加 more，tall 是单音节词直接加 -er。"
    )

    # 判断题
    add_question(doc, 11,
        "'Apple' is a countable noun.",
        ["A. 正确", "B. 错误"],
        "正确",
        "'apple' 是可数名词，可以说 an apple, two apples。"
    )

    add_question(doc, 12,
        "The past tense of 'read' is 'readed'.",
        ["A. 正确", "B. 错误"],
        "错误",
        "'read' 的过去式仍然是 'read'（拼写不变，但发音变为/red/），这是不规则动词。"
    )

    doc.save(os.path.join(OUTPUT_DIR, "英语基础题库.docx"))
    print(f"✓ 英语基础题库.docx (12题: 10单选+2判断)")


def create_math_basics():
    """数学基础题库"""
    doc = Document()

    add_question(doc, 1,
        "2的10次方等于多少？",
        ["A. 512", "B. 1000", "C. 1024", "D. 2048"],
        "C",
        "2^10 = 1024，这是计算机科学中非常重要的数字。"
    )

    add_question(doc, 2,
        "一个三角形三个内角之和是多少度？",
        ["A. 90度", "B. 180度", "C. 270度", "D. 360度"],
        "B",
        "三角形内角和恒等于180度，这是欧几里得几何的基本定理。"
    )

    add_question(doc, 3,
        "圆的面积公式是？",
        ["A. 2πr", "B. πr²", "C. πd", "D. 4πr²"],
        "B",
        "圆的面积 = π × r²，其中 r 是半径。2πr 是周长公式。"
    )

    add_question(doc, 4,
        "以下哪个数是质数？",
        ["A. 9", "B. 15", "C. 17", "D. 21"],
        "C",
        "17只能被1和自身整除，是质数。9=3×3，15=3×5，21=3×7，都不是质数。"
    )

    add_question(doc, 5,
        "1+2+3+...+100 等于多少？",
        ["A. 5000", "B. 5050", "C. 5100", "D. 10000"],
        "B",
        "等差数列求和公式：S = n(n+1)/2 = 100×101/2 = 5050。这是高斯小时候的经典题目。"
    )

    add_question(doc, 6,
        "一个正方体有几个面？",
        ["A. 4个", "B. 6个", "C. 8个", "D. 12个"],
        "B",
        "正方体有6个面、12条棱、8个顶点。"
    )

    add_question(doc, 7,
        "√144 等于多少？",
        ["A. 10", "B. 11", "C. 12", "D. 14"],
        "C",
        "12 × 12 = 144，所以 √144 = 12。"
    )

    add_question(doc, 8,
        "0.1 + 0.2 在计算机中等于多少？",
        ["A. 0.3", "B. 0.30000000000000004", "C. 0.28", "D. 0.31"],
        "B",
        "这是浮点数精度问题。由于二进制无法精确表示0.1和0.2，所以0.1+0.2的结果不是精确的0.3。"
    )

    add_question(doc, 9,
        "以下哪个等式成立？",
        ["A. (a+b)² = a² + b²", "B. (a+b)² = a² + 2ab + b²", "C. (a+b)² = a² - 2ab + b²", "D. (a+b)² = a² + ab + b²"],
        "B",
        "完全平方公式：(a+b)² = a² + 2ab + b²。选项A是常见错误。"
    )

    add_question(doc, 10,
        "1千米等于多少米？",
        ["A. 100米", "B. 500米", "C. 1000米", "D. 10000米"],
        "C",
        "1千米(km) = 1000米(m)，这是国际单位制中的长度换算。"
    )

    # 判断题
    add_question(doc, 11,
        "所有的正方形都是长方形。",
        ["A. 正确", "B. 错误"],
        "正确",
        "正方形是特殊的长方形（四条边都相等的长方形），所以所有正方形都是长方形。"
    )

    add_question(doc, 12,
        "0是最小的自然数。",
        ["A. 正确", "B. 错误"],
        "正确",
        "在我国数学标准中，自然数从0开始：0, 1, 2, 3, ... 所以0是最小的自然数。"
    )

    add_question(doc, 13,
        "π（圆周率）是一个有理数。",
        ["A. 正确", "B. 错误"],
        "错误",
        "π是一个无理数（无限不循环小数），约等于3.14159...，不能表示为两个整数的比。"
    )

    add_question(doc, 14,
        "两条平行线永远不会相交。",
        ["A. 正确", "B. 错误"],
        "正确",
        "在欧几里得几何中，平行线的定义就是在同一平面内永不相交的两条直线。"
    )

    add_question(doc, 15,
        "-5 比 -3 大。",
        ["A. 正确", "B. 错误"],
        "错误",
        "在数轴上，-5 在 -3 的左边，所以 -5 < -3。负数中，绝对值越大的数越小。"
    )

    doc.save(os.path.join(OUTPUT_DIR, "数学基础题库.docx"))
    print(f"✓ 数学基础题库.docx (15题: 10单选+5判断)")


if __name__ == "__main__":
    create_computer_basics()
    create_english_basics()
    create_math_basics()
    print(f"\n所有题库已生成到: {OUTPUT_DIR}")
